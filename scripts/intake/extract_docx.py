"""Extract raw text from a .docx. Usage: python extract_docx.py <path.docx>"""
import sys
from docx import Document


def extract(path):
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("Usage: python extract_docx.py <path.docx>", file=sys.stderr)
        sys.exit(1)
    print(extract(sys.argv[1]))
